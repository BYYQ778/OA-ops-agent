"""
生成项目答辩 Word 文档 — 知识问答 + 知识库模块技术详解 + 豆包 PPT 提示词
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# ============ 封面 ============
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('OA运维多智能Agent巡检问答系统')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('知识问答 & 知识库模块 · 技术实现详解')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()
version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('v2.4.1  |  2026年7月  |  项目答辩材料')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x8B, 0x8F, 0xA3)

doc.add_page_break()

# ============ 目录页 ============
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述与技术架构',
    '二、知识问答模块 — 技术实现',
    '    2.1 整体架构',
    '    2.2 Agentic RAG（LangGraph ReAct 多步推理）',
    '    2.3 多轮对话（Chatbot）',
    '    2.4 双 LLM 后端兼容（Ollama / DeepSeek）',
    '    2.5 批量问答',
    '    2.6 前端实现',
    '三、知识库模块 — 技术实现',
    '    3.1 文档解析管线',
    '    3.2 向量存储与检索（ChromaDB）',
    '    3.3 知识图谱（NetworkX）',
    '    3.4 实体提取（双策略）',
    '    3.5 图谱构建与探索',
    '    3.6 前端实现',
    '四、关键技术决策与创新点',
    '五、技术栈总览',
    '六、API 接口清单',
    '附录：豆包 AI 自动生成答辩 PPT 提示词',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============ 一、项目概述 ============
doc.add_heading('一、项目概述与技术架构', level=1)

doc.add_paragraph(
    '本项目是一个面向企业OA运维场景的智能Agent系统，支持实时监控、自动巡检、日志分析、'
    '知识问答、知识库管理等8大功能模块。本文档聚焦于"知识问答"和"知识库"两个核心智能模块，'
    '详细阐述其技术架构、实现原理和创新设计。'
)

doc.add_heading('系统整体架构', level=2)
doc.add_paragraph(
    '后端：FastAPI（47个API端点）+ LangChain + LangGraph + ChromaDB + NetworkX + SQLite + Ollama/DeepSeek\n'
    '前端：纯 HTML/CSS/JS（8页面）+ Chart.js 4.4.0 + EventSource SSE\n'
    'Agent模块：11个Agent + 11个工具模块\n'
    '部署方式：Python 虚拟环境 + Uvicorn，双击 .bat 一键启动'
)

doc.add_heading('知识问答与知识库的关系', level=2)
doc.add_paragraph(
    '两个模块紧密协作：知识库负责文档解析、向量化存储和知识图谱构建；'
    '知识问答模块基于知识库的检索结果，通过 Agentic RAG（LangGraph ReAct）'
    '进行多步推理，生成精准回答。两者共同构成"文档入库 → 智能检索 → 多步推理 → 精准回答"'
    '的完整知识管理闭环。'
)

doc.add_page_break()

# ============ 二、知识问答模块 ============
doc.add_heading('二、知识问答模块 — 技术实现', level=1)

# 2.1
doc.add_heading('2.1 整体架构', level=2)
doc.add_paragraph(
    '知识问答模块是整个系统的智能核心，采用 Agentic RAG（检索增强生成）架构，'
    '由以下核心组件构成：'
)

# 架构表格
table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.autofit = True
headers = ['组件', '技术选型', '职责']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

data = [
    ['LLM 大模型', 'DeepSeek-Chat / Qwen3:8b (Ollama)', '文本生成、多步推理、路由决策'],
    ['嵌入模型', 'paraphrase-multilingual-MiniLM-L12-v2', '文档向量化（本地 CPU 运行，118MB）'],
    ['向量数据库', 'ChromaDB（LangChain 集成）', '语义检索、相似度匹配'],
    ['Agent 框架', 'LangGraph StateGraph', 'ReAct 多步推理流程编排'],
    ['对话记忆', '内存 Dict + FIFO 淘汰', '多轮对话上下文管理'],
    ['前端', '纯 HTML/CSS/JS', 'Chatbot UI、Markdown 渲染、命令推荐'],
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# 2.2
doc.add_heading('2.2 Agentic RAG — LangGraph ReAct 多步推理', level=2)

doc.add_paragraph(
    '这是本系统最核心的技术创新。传统 RAG 是"检索一次 → 回答一次"的单步模式，'
    '面对复杂运维问题（如"数据库连接失败，可能的原因有哪些？如何逐一排查？"）往往力不从心。'
    '本系统实现了基于 LangGraph 的 Agentic RAG，支持多轮检索-推理循环。'
)

doc.add_heading('StateGraph 状态图结构', level=3)
doc.add_paragraph(
    '定义 AgenticState（TypedDict），包含7个状态字段：\n'
    '  • messages：累计对话消息\n'
    '  • question：用户当前问题\n'
    '  • kb_context：向量检索累积结果\n'
    '  • kg_context：知识图谱检索累积结果\n'
    '  • reasoning_steps：推理步数计数器\n'
    '  • next_action：路由决策（search_kb / search_kg / explore_graph / answer）\n'
    '  • final_answer：最终回答'
)

doc.add_heading('五个处理节点', level=3)

nodes_data = [
    ('router（路由器）',
     'LLM 根据当前已获取的信息和可用工具，自主决策下一步操作。'
     '通过文本 Prompt 指导 LLM 输出 "NEXT_ACTION: <tool>" 格式，'
     '用正则解析路由目标。最大推理步数为5步，超出后强制进入 answer。'),
    ('search_kb（向量检索）',
     '调用 ChromaDB similarity_search()，以用户问题或子问题为查询，'
     '在知识库中检索 top_k=5 的最相关文档片段。结果追加到 kb_context。'),
    ('search_kg（图谱检索）',
     '在知识图谱中搜索匹配的实体节点，获取其描述、关联实体和来源文档。'
     '结果追加到 kg_context。'),
    ('explore_graph（图谱探索）',
     '以检索到的 Top3 实体为起点，BFS 提取 2 跳子图，'
     '发现间接关联的技术、组织、人员等实体。结果追加到 kg_context。'),
    ('answer（生成回答）',
     '将 kb_context + kg_context 拼接为最终 Prompt，'
     '由 LLM 严格基于检索内容生成回答。附上文档来源引用。'),
]
for name, desc in nodes_data:
    p = doc.add_paragraph()
    run = p.add_run(f'● {name}：')
    run.font.bold = True
    p.add_run(desc)

doc.add_heading('路由循环机制', level=3)
doc.add_paragraph(
    'Entry → router → {search_kb | search_kg | explore_graph} → router → ... → answer → END\n\n'
    '这是一个典型的 ReAct (Reasoning + Acting) 模式：LLM 先思考需要什么信息，'
    '然后调用相应工具获取，根据新信息再次思考，直到信息充分后生成最终回答。'
    '整个流程由 LangGraph 的 StateGraph 编译为可执行的 agentic_graph，通过 invoke() 驱动。'
)

doc.add_heading('Ollama 兼容策略', level=3)
doc.add_paragraph(
    '核心挑战：Ollama 本地模型（qwen3:8b）不支持原生的 Function Calling / Tool Calling API。\n\n'
    '解决方案：完全绕过工具调用 API，使用"文本 Prompt + 正则解析"的路由机制。\n'
    '  • Router Prompt 中明确要求 LLM 输出 "NEXT_ACTION: <action>" 格式\n'
    '  • 后端用正则 r"NEXT_ACTION:\\s*(\\w+)" 解析 LLM 响应\n'
    '  • 验证 action 是否在合法工具列表中\n'
    '  • 如果 LLM 未按要求输出，自动降级：内容超100字符无 NEXT_ACTION → 视为 answer；'
    'LLM 调用异常 → 默认 search_kb\n\n'
    '这一设计使得 Ollama 本地模型和 DeepSeek 云端模型共享同一套 Agentic RAG 代码路径，'
    '无需维护两套实现。'
)

doc.add_heading('降级与容错', level=3)
doc.add_paragraph(
    '系统设计了多层降级机制，确保在任何异常情况下仍能提供基本服务：\n'
    '  1. LangGraph 异常 → 自动回退到简单 RAG（检索 + 单次 LLM 调用）\n'
    '  2. LLM 不可用 → 返回原始检索结果（[检索降级模式]），用户可自行参考\n'
    '  3. KG 不可用 → search_kg/explore_graph 自动降级为 search_kb\n'
    '  4. Agentic RAG 关闭 → 完全回退到 v2.3 版本的简单 RAG\n'
    '  5. MinerU 不可用 → 回退到传统文档解析器'
)

# 2.3
doc.add_heading('2.3 多轮对话（Chatbot）', level=2)

doc.add_paragraph(
    '系统支持完整的上下文感知多轮对话，而非每次独立的单轮问答。\n\n'
    '对话记忆机制：\n'
    '  • 存储：内存 Dict，key=conversation_id，value=消息列表\n'
    '  • 容量：最多保留 chat_max_history=30 轮（60条消息），超出后 FIFO 淘汰旧消息\n'
    '  • 上下文窗口：每次传给 LLM 最近 chat_max_turns=10 轮对话\n'
    '  • 消息截断：每条消息截断为 500 字，避免 Token 超限\n'
    '  • 格式化：最近消息在前，旧消息在后；用户和 AI 的角色标签明确区分\n\n'
    '交互流程：\n'
    '  用户输入 → 格式化历史 → Agentic RAG 推理 → 生成回答 → 保存对话 → 返回结果\n\n'
    '对话管理 API：\n'
    '  POST /api/kb/chat — 发送消息（自动关联历史）\n'
    '  POST /api/kb/chat/clear — 清除指定对话\n'
    '  GET /api/kb/chat/history — 获取历史消息'
)

# 2.4
doc.add_heading('2.4 双 LLM 后端兼容', level=2)

doc.add_paragraph(
    '系统支持一键切换 LLM 后端，满足不同场景需求：\n\n'
    '  本机离线模式（Ollama + qwen3:8b）：\n'
    '    • 零网络依赖，数据不出企业内网\n'
    '    • 使用可配置的 Ollama 端点（默认 http://localhost:11434/v1）\n'
    '    • 无需 API Key\n\n'
    '  云端在线模式（DeepSeek API）：\n'
    '    • 更强的推理能力和更快的响应速度\n'
    '    • 支持 Function Calling（实体提取时使用）\n'
    '    • 通过 .env 文件管理 API Key，不暴露在配置文件中\n\n'
    '切换机制：\n'
    '  • 前端：系统设置页面 → 选择模式 → 填写参数 → 保存\n'
    '  • 后端：POST /api/config/save 写入 config.yaml → 热重载 → 重置 KB Agent\n'
    '  • 无需重启服务，下次问答即刻生效'
)

# 2.5
doc.add_heading('2.5 批量问答', level=2)
doc.add_paragraph(
    '支持一次性提交最多20个问题，系统并行处理后返回全部答案。\n\n'
    '技术实现：\n'
    '  • 输入：每行一个问题，换行分隔\n'
    '  • 并行度：ThreadPoolExecutor，max_workers=5\n'
    '  • 限流：最多 batch_max_questions=20 题\n'
    '  • 容错：单个问题失败不影响其他问题\n'
    '  • 前端：可折叠的 Q&A 卡片，点击展开查看答案'
)

# 2.6
doc.add_heading('2.6 前端实现', level=2)
doc.add_paragraph(
    'Chatbot UI 特性：\n'
    '  • 消息气泡：用户靠右蓝底白字，AI 靠左灰底黑字，最大宽度75%\n'
    '  • Markdown 渲染：支持代码块（```）和行内代码（`）的格式化显示\n'
    '  • 思考状态：发送后显示斜体"思考中..."占位，收到回复后替换\n'
    '  • 命令推荐：每次回复后自动匹配运维命令大全中的相关命令，点击可快速追问\n'
    '  • 新对话：一键清除会话历史，开始全新对话\n'
    '  • 键盘操作：Enter 发送，无需点击按钮\n\n'
    '通信方式：\n'
    '  • 聊天：HTTP POST（FormData），同步请求-响应模式\n'
    '  • 仪表盘：SSE（EventSource）实时推送，30秒心跳保活\n'
    '  • 所有 API 返回 JSON 格式'
)

doc.add_page_break()

# ============ 三、知识库模块 ============
doc.add_heading('三、知识库模块 — 技术实现', level=1)

# 3.1
doc.add_heading('3.1 文档解析管线', level=2)

doc.add_paragraph(
    '系统支持多种格式文档的自动解析，从原始文件到可检索的结构化文本，全流程自动化。'
)

doc.add_heading('支持的文档格式', level=3)
table2 = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table2.autofit = True
headers2 = ['格式', '解析库', '处理方式']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

fmt_data = [
    ['PDF', 'PyPDF2', '逐页提取文本，自动检测扫描件'],
    ['DOCX', 'python-docx', '段落迭代提取，跳过空段落'],
    ['PPTX', 'python-pptx', '逐幻灯片提取，含文本框和表格'],
    ['HTML', 'BeautifulSoup + lxml', '剥离脚本/样式/导航，提取正文'],
    ['TXT', '内置 open()', '自动编码检测（utf-8→gbk→gb2312→latin-1）'],
    ['图片', 'CnOCR', '中文 OCR 识别（JPG/PNG/BMP/TIFF/WEBP）'],
]
for i, row_data in enumerate(fmt_data):
    for j, val in enumerate(row_data):
        table2.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('文本处理管线', level=3)
doc.add_paragraph(
    '原始文件 → 格式解析 → clean_text()（去除控制字符、规范化空白）→ '
    'split_text()（智能分块）→ LangChain Document 对象 → ChromaDB 向量化存储\n\n'
    '智能分块策略（split_text）：\n'
    '  • chunk_size=500 字符，chunk_overlap=50 字符\n'
    '  • 中文化优化：优先以"。"或换行符作为分割点\n'
    '  • 边界检测：在窗口50%位置之后查找自然断点，避免截断句子\n\n'
    '去重机制：\n'
    '  • 导入前计算文件的 MD5 哈希\n'
    '  • 查询 ChromaDB（where={"file_hash": md5}）判断是否已存在\n'
    '  • 已存在的文档自动跳过，避免重复导入'
)

doc.add_heading('MinerU 多模态解析（可选）', level=3)
doc.add_paragraph(
    '系统集成了 MinerU（magic-pdf）作为增强解析器，用于处理复杂排版的 PDF：\n'
    '  • 自动分类文档类型（OCR 模式 vs 文本模式）\n'
    '  • 提取 Markdown 格式的表格和 LaTeX 格式的公式\n'
    '  • 扫描件 PDF 自动路由到 OCR 管线\n'
    '  • 懒加载设计：仅在 config 中启用且文件需要时才导入（~3GB 模型）\n'
    '  • 默认关闭，失败自动回退传统解析器'
)

# 3.2
doc.add_heading('3.2 向量存储与检索（ChromaDB）', level=2)

doc.add_paragraph(
    '向量化流程：\n'
    '  1. 文档分块后，每块创建 LangChain Document 对象（含 metadata）\n'
    '  2. 通过 HuggingFaceEmbeddings 将文本转为 384 维向量\n'
    '  3. 向量 + 元数据存入 ChromaDB（持久化到 data/chroma_db/）\n'
    '  4. 检索时使用余弦相似度（similarity_search）返回 top_k=5 最相关片段\n\n'
    '嵌入模型特点：\n'
    '  • 模型：paraphrase-multilingual-MiniLM-L12-v2（118MB）\n'
    '  • 纯本地 CPU 运行，无需 GPU\n'
    '  • 支持中英文，通过 HuggingFace 镜像（hf-mirror.com）下载\n'
    '  • 两阶段加载：优先使用本地缓存（local_files_only=True），失败后在线下载\n'
    '  • normalize_embeddings=True，确保向量可比性\n\n'
    '元数据结构：\n'
    '  {source, file_path, chunk_index, total_chunks, file_hash, import_time, chunk_size}'
)

# 3.3
doc.add_heading('3.3 知识图谱（NetworkX）', level=2)

doc.add_paragraph(
    '系统构建了一个基于 NetworkX 的知识图谱，用于捕捉文档中实体间的关联关系，'
    '弥补纯向量检索在语义关联发现上的不足。'
)

doc.add_heading('技术选型', level=3)
doc.add_paragraph(
    '  • 图引擎：NetworkX DiGraph（有向图）\n'
    '  • 持久化：JSONL 文件（nodes.jsonl + edges.jsonl），每行一条 JSON 记录\n'
    '  • 原子写入：先写 .tmp 文件，再 os.replace()，防断电损坏\n'
    '  • 零外部依赖：不依赖 Neo4j 等图数据库，保持部署简便\n\n'
    '设计理念：\n'
    '  "KG 只是 Chroma 的补充，而非替代。"\n'
    '  通过共现关系自动建图，无需人工定义 Schema。\n'
    '  KG 故障不影响向量检索和基本问答功能。'
)

doc.add_heading('实体类型（5种）', level=3)
doc.add_paragraph(
    '  • TECHNOLOGY：技术/软件/协议（如 Nginx、MySQL、HTTPS）\n'
    '  • ORGANIZATION：组织/部门/厂商（如 信息中心、运维部）\n'
    '  • PERSON：人员姓名\n'
    '  • LOCATION：地点/机房位置\n'
    '  • CONCEPT：抽象概念/业务术语（如 审批流程、灾备方案）'
)

doc.add_heading('共现建边策略', level=3)
doc.add_paragraph(
    '同一个文档分块内出现的任意两个实体之间自动建立 co_occurs 边。\n'
    '  • 权重累计：多次共现时边权重递增\n'
    '  • 全连接：同一 chunk 内所有实体两两建边（而非仅相邻实体）\n'
    '  • 去重合并：同一实体跨文档出现时合并来源信息，frequency 递增'
)

doc.add_heading('图谱操作', level=3)
ops_data = [
    ('实体搜索', '按名称/描述模糊匹配（不区分大小写），支持按类型筛选，最多返回20条'),
    ('邻居探索', '返回指定节点的所有邻居，按边权重降序排列'),
    ('子图提取', 'BFS 遍历，提取以指定节点为中心的 N 跳子图（默认 depth=2）'),
    ('路径发现', '基于 NetworkX shortest_path() 查找两实体间的最短关联路径'),
    ('文档删除同步', '删除文档时同步清理 KG 中的孤立实体和边'),
]
for name, desc in ops_data:
    p = doc.add_paragraph()
    run = p.add_run(f'● {name}：')
    run.font.bold = True
    p.add_run(desc)

# 3.4
doc.add_heading('3.4 实体提取（双策略）', level=2)

doc.add_paragraph(
    '实体提取是整个 KG 构建的入口环节。系统针对不同 LLM 后端实现了两种提取策略，'
    '统一通过 EntityExtractor 类对外暴露 extract() 接口。'
)

doc.add_heading('策略一：Function Calling（DeepSeek/OpenAI）', level=3)
doc.add_paragraph(
    '  • 定义 extract_entities 工具，entity_type 枚举5种类型\n'
    '  • 使用 tool_choice 强制 LLM 调用该工具\n'
    '  • 解析 response.tool_calls 获取结构化实体列表\n'
    '  • 失败时自动回退到策略二'
)

doc.add_heading('策略二：Prompt Engineering（Ollama/qwen3:8b）', level=3)
doc.add_paragraph(
    '  • 构造 System Prompt 详细描述5种实体类型及示例\n'
    '  • 要求 LLM 输出严格 JSON 数组格式（```json ... ```）\n'
    '  • 三层解析降级：\n'
    '    A. json.loads() 直接解析\n'
    '    B. 正则提取代码块中的 JSON\n'
    '    C. 逐行正则匹配 "类型: 名称" 模式（兜底）\n'
    '  • 长文本自动分块（每块 ≤ 1200 字符，50 字符重叠）'
)

doc.add_heading('质量控制', level=3)
doc.add_paragraph(
    '  • 类型校验：必须在5种 ENTITY_TYPES 中\n'
    '  • 名称校验：非空且长度 < 100 字符\n'
    '  • 去重合并：按 (类型, 名称) 去重，后者补充前者的描述'
)

# 3.5
doc.add_heading('3.5 图谱构建与探索', level=2)

doc.add_paragraph(
    '完整构建流程：\n'
    '  文档导入 → 文本解析 → 分块（800字/块）→ 逐块实体提取 → '
    '实体去重合并 → 建节点（upsert）→ 块内共现建边 → 增量持久化（JSONL）\n\n'
    '关键设计：\n'
    '  • 节点 ID 稳定：md5("{type}:{name}")[:12]，跨次构建保持不变\n'
    '  • 增量更新：新文档追加到已有图谱，不会全量重建\n'
    '  • Frequency 计数：记录实体在不同文档和分块中的出现频次\n'
    '  • Sources 追溯：每个实体记录其来源文档，支持按文档删除\n\n'
    '探索接口：\n'
    '  POST /api/kg/search — 模糊搜索实体\n'
    '  POST /api/kg/explore — 提取子图（含邻居信息）\n'
    '  POST /api/kg/path — 最短路径查询\n'
    '  GET /api/kg/stats — 图谱统计（节点数、边数、类型分布）'
)

# 3.6
doc.add_heading('3.6 前端知识库界面', level=2)
doc.add_paragraph(
    '知识库页面分为两个子标签页：\n\n'
    '  文档管理：\n'
    '    • 拖拽上传（支持7种格式）\n'
    '    • 文档列表查看（显示分块数、字符数、导入时间）\n'
    '    • 单个删除 / 一键清空（含确认对话框）\n'
    '    • 知识库统计（总文档数、总分块数、KG 节点/边数）\n\n'
    '  图谱探索：\n'
    '    • 实体搜索（关键词 + 类型筛选）\n'
    '    • 子图探索（点击实体查看2跳邻居）\n'
    '    • 路径发现（输入源和目标实体查找关联路径）\n'
    '    • 图谱统计（自动加载）\n'
    '    • 纯 HTML/CSS 渲染（无 D3.js 等外部依赖），类型标签颜色区分\n\n'
    '  力导向图可视化（v2.4.1 新增）：\n'
    '    • 技术选型：vis.js (vis-network)，CDN 引入 ~200KB，零构建步骤\n'
    '    • 节点着色：5种实体类型映射5种颜色（蓝/橙/紫/绿/灰），浅色背景配深色字，深色配白字\n'
    '    • 节点大小：按频次动态缩放（25~55px），高频实体更突出\n'
    '    • 字体描边：半透明白色投影确保在任何背景色上都可读\n'
    '    • 边样式：曲线箭头，宽度按权重递增（1~5px），hover 显示关系详情\n'
    '    • 交互：拖拽节点调整布局、滚轮缩放、点击节点以该节点为中心展开2跳邻居\n'
    '    • 导航：历史栈（20层）+ "↩ 返回"按钮，支持逐级回退\n'
    '    • 双视图切换：力导向图 ⇄ 卡片列表一键切换\n'
    '    • 全图模式：聚合 Top10 实体子图展示完整知识图谱'
)

doc.add_page_break()

# ============ 四、关键技术决策 ============
doc.add_heading('四、关键技术决策与创新点', level=1)

innovations = [
    ('1. Agentic RAG 替代简单 RAG',
     '传统 RAG 是单步"检索→回答"，无法处理需要多步推理的复杂问题。'
     '本系统使用 LangGraph ReAct 模式，LLM 自主决定需要检索什么、是否需要进一步探索图谱，'
     '最多5步推理循环，显著提升复杂运维问题的回答质量。'),

    ('2. 文本路由替代 Function Calling',
     '由于 Ollama 本地模型不支持原生 Tool Calling，系统设计了一套"Prompt + 正则解析"的通用路由机制。'
     '核心思路：让 LLM 在 Response 中输出 "NEXT_ACTION: <tool>" 格式，'
     '后端用正则解析后执行相应工具，将结果注入下一轮对话。'
     '这种设计统一了 Ollama 和 DeepSeek 的代码路径，无需维护两套实现。'),

    ('3. 向量检索 + 知识图谱双引擎',
     'ChromaDB 负责语义相似度匹配，知识图谱负责结构化关联发现。'
     '两者互为补充：向量检索擅长"这段文字在讲什么"，图谱擅长"这个东西和哪些东西有关系"。'
     '在回答"OA 审批流程依赖哪些系统？"这类问题时，图谱能直接输出关联节点，'
     '无需在文档片段中逐段查找。'),

    ('4. 全链路降级设计',
     '每一层都可能失败，每一层都有降级方案：\n'
     '  Agentic RAG → 简单 RAG → 纯检索\n'
     '  KG 不可用 → 仅向量检索\n'
     '  MinerU 不可用 → 传统解析器\n'
     '  嵌入模型离线 → 在线下载\n'
     '  确保系统在任何异常情况下仍能提供基本服务。'),

    ('5. 零外部服务依赖的 KG',
     '不依赖 Neo4j、MySQL 等外部数据库，使用 NetworkX 内存图 + JSONL 文件持久化。'
     '部署时无需额外安装任何数据库，一个 Python 虚拟环境即可运行完整系统。'),

    ('6. 中文化优化的分块策略',
     '文本分块时优先在中文句号"。"处切割，而非机械地按字符数截断。'
     '避免了"服务器需要连接到..."和"...数据库才能正常工作"被切成两个 chunk 的问题，'
     '显著提升检索召回质量。'),

    ('7. 本地化部署友好',
     '所有模型均可本地运行：嵌入模型（118MB CPU）、Ollama + qwen3:8b。'
     '通过 HuggingFace 镜像（hf-mirror.com）加速模型下载。'
     '企业内网环境下零外部 API 依赖即可运行完整知识问答功能。'),

    ('8. 知识图谱力导向可视化（v2.4.1）',
     '集成 vis.js 力导向图引擎（~200KB CDN），零构建步骤兼容原生 JS 架构。'
     '实现节点按类型着色/频次定大小/边按权重定粗细、拖拽缩放、点击展开邻居、'
     '历史栈回退等完整交互。无需 D3.js 等重型库，纯 HTML 容器渲染，与现有系统无缝集成。'),
]

for title, desc in innovations:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============ 五、技术栈总览 ============
doc.add_heading('五、技术栈总览', level=1)

stack_table = doc.add_table(rows=18, cols=3, style='Light Grid Accent 1')
stack_table.autofit = True
sh = ['层级', '技术', '用途']
for i, h in enumerate(sh):
    stack_table.rows[0].cells[i].text = h
    for p in stack_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

stack_data = [
    ['LLM 框架', 'LangChain 1.0+ / LangGraph 0.2+', 'RAG 管线、Agent 编排、ReAct 循环'],
    ['LLM 后端', 'DeepSeek API / Ollama (qwen3:8b)', '文本生成、推理、路由决策'],
    ['嵌入模型', 'sentence-transformers (MiniLM-L12-v2)', '文档向量化、语义检索'],
    ['向量数据库', 'ChromaDB 1.0+（LangChain 集成）', '持久化向量存储、相似度搜索'],
    ['图引擎', 'NetworkX 3.2+', '知识图谱构建、BFS 子图、最短路径'],
    ['Web 框架', 'FastAPI 0.115+ / Uvicorn', 'REST API 服务'],
    ['前端', '原生 HTML/CSS/JS + Chart.js 4.4.0', '8页面 SPA、SSE 实时推送'],
    ['文档解析', 'PyPDF2 / python-docx / python-pptx / BS4', '多格式文档提取'],
    ['OCR', 'CnOCR 2.3+', '图片文字识别'],
    ['MinerU', 'magic-pdf（可选，~3GB）', '复杂 PDF 多模态解析'],
    ['LLM API', 'OpenAI SDK 2.0+ (ChatOpenAI)', '统一 LLM 调用接口'],
    ['数据库', 'SQLite (APScheduler 任务)', '巡检记录持久化'],
    ['配置', 'PyYAML 6.0 + .env', '配置管理、环境变量注入'],
    ['模板', 'Jinja2 3.1+', 'HTML 模板渲染'],
    ['调度', 'APScheduler 3.10+', '定时巡检任务'],
    ['NLP', 'scikit-learn / numpy', '文本处理、向量计算'],
    ['分词', 'tiktoken', 'Token 计数'],
]
for i, row_data in enumerate(stack_data):
    for j, val in enumerate(row_data):
        stack_table.rows[i+1].cells[j].text = val

doc.add_page_break()

# ============ 六、API 接口清单 ============
doc.add_heading('六、API 接口清单（知识问答 & 知识库相关）', level=1)

api_table = doc.add_table(rows=15, cols=3, style='Light Grid Accent 1')
api_table.autofit = True
ah = ['端点', '方法', '功能']
for i, h in enumerate(ah):
    api_table.rows[0].cells[i].text = h
    for p in api_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

api_data = [
    ['/api/kb/chat', 'POST', '多轮对话（核心接口）'],
    ['/api/kb/chat/clear', 'POST', '清除对话历史'],
    ['/api/kb/chat/history', 'GET', '获取对话历史'],
    ['/api/kb/ask', 'POST', '单次问答（无对话记忆）'],
    ['/api/kb/import', 'POST', '上传文档导入知识库'],
    ['/api/kb/list', 'GET', '列出已导入文档'],
    ['/api/kb/stats', 'GET', '知识库统计（含KG）'],
    ['/api/kb/delete', 'POST', '删除指定文档'],
    ['/api/kb/clear', 'POST', '清空全部知识库'],
    ['/api/kb/batch-ask', 'POST', '批量问答（最多20题）'],
    ['/api/kg/search', 'POST', '图谱实体搜索'],
    ['/api/kg/explore', 'POST', '图谱子图探索'],
    ['/api/kg/path', 'POST', '两实体间最短路径'],
    ['/api/kg/stats', 'GET', '图谱统计信息'],
]
for i, row_data in enumerate(api_data):
    for j, val in enumerate(row_data):
        api_table.rows[i+1].cells[j].text = val

doc.add_page_break()

# ============ 附录：豆包 PPT 提示词 ============
doc.add_heading('附录：豆包 AI 自动生成答辩 PPT 提示词', level=1)

doc.add_paragraph(
    '以下提示词可直接复制粘贴到豆包（Doubao）AI 中，自动生成项目答辩 PPT。'
    '建议分两次使用：第一次生成整体框架，第二次针对具体页面补充细节。'
)

doc.add_heading('提示词一：生成完整答辩 PPT（12-15页）', level=2)

prompt1 = doc.add_paragraph()
prompt1.paragraph_format.space_before = Pt(8)
run = prompt1.add_run(
    '请帮我生成一份项目答辩 PPT，主题是《OA运维多智能Agent巡检问答系统 — '
    '知识问答与知识库模块技术实现》，要求12-15页。\n\n'
    '=== PPT 内容要求 ===\n\n'
    '第1页-封面：标题"OA运维多智能Agent巡检问答系统"，副标题"知识问答 & 知识库模块技术实现"，'
    'v2.4版本，2026年7月\n\n'
    '第2页-项目背景：企业OA运维面临的问题（文档分散、故障排查依赖经验、新人上手慢），'
    '本系统的解决思路（智能知识管理 + AI 问答）\n\n'
    '第3页-系统架构总览：FastAPI + LangChain + LangGraph + ChromaDB + NetworkX + Ollama，'
    '8大功能模块，前端8页面\n\n'
    '第4页-知识问答模块架构：Agentic RAG（LangGraph ReAct）、多轮对话、双LLM后端\n\n'
    '第5页-Agentic RAG 核心原理：StateGraph 5节点（router/search_kb/search_kg/explore_graph/answer），'
    'ReAct 循环（最多5步推理），文本路由机制\n\n'
    '第6页-Ollama/DeepSeek 双后端兼容：核心挑战（Ollama不支持Function Calling），'
    '解决方案（Prompt+正则解析的通用路由），一键热切换\n\n'
    '第7页-多轮对话实现：内存对话记忆（30轮FIFO），上下文窗口（10轮），消息截断\n\n'
    '第8页-知识库模块架构：文档解析→分块→向量化→ChromaDB，'
    '实体提取→知识图谱（NetworkX），双引擎检索\n\n'
    '第9页-文档解析管线：支持7种格式（PDF/DOCX/PPTX/HTML/TXT/图片OCR/MinerU），'
    '智能中文化分块，MD5去重\n\n'
    '第10页-知识图谱：5种实体类型，共现建边，JSONL持久化，BFS子图探索，最短路径\n\n'
    '第11页-力导向图可视化（v2.4.1）：vis.js引擎，节点按类型着色/频次定大小，'
    '拖拽缩放，点击展开邻居，历史回退，双视图切换\n\n'
    '第12页-实体提取双策略：DeepSeek用Function Calling，Ollama用Prompt工程+三层解析降级\n\n'
    '第13页-降级设计：每一层都有 fallback（Agentic RAG→简单RAG→纯检索，'
    'KG不可用→仅向量，MinerU不可用→传统解析）\n\n'
    '第14页-技术栈与API清单：LangChain/LangGraph/ChromaDB/NetworkX/FastAPI/vis.js，15个KB相关API\n\n'
    '第15页-创新点总结：Agentic RAG多步推理、文本路由统一双后端、向量+图谱双引擎、'
    '全链路降级、零外部依赖KG、中文化分块、力导向图可视化\n\n'
    '第16页-演示与展望：当前v2.4.1已实现的功能，未来方向（流式输出、对话持久化、'
    '多模态问答、权限管理）\n\n'
    '=== 格式要求 ===\n'
    '1. 每页PPT包含：标题 + 3-5个要点（bullet points） + 1句总结\n'
    '2. 技术架构页建议用流程图或架构图描述（用文字描述即可，豆包会生成）\n'
    '3. 配色建议：主色 #4F46E5（靛蓝），辅色 #3B82F6（蓝）\n'
    '4. 风格：专业技术风，简洁明了\n'
    '5. 每页底部加页码'
)

doc.add_heading('提示词二：生成技术架构图描述', level=2)

prompt2 = doc.add_paragraph()
run = prompt2.add_run(
    '请帮我生成一张技术架构图的文字描述，用于 PPT 中的架构页。\n\n'
    '架构包含以下层次（从上到下）：\n'
    '第1层-用户界面：浏览器（8个页面：实时监控、经典巡检、日志分析、知识问答、知识库、'
    '诊断工具箱、命令大全、系统设置）\n'
    '第2层-API网关：FastAPI + Uvicorn（47个端点，FormData/JSON/SSE）\n'
    '第3层-Agent层：KnowledgeBaseAgent（RAG问答、多轮对话）| EntityExtractor（实体提取）| '
    'KG Builder（图谱构建）| AI Reporter（报告生成）| 巡检Agent（6个诊断模块）\n'
    '第4层-工具层：ChromaDB（向量存储）| NetworkX（知识图谱）| CnOCR（文字识别）| '
    'PyPDF2/DOCX/PPTX/BS4（文档解析）| MinerU（多模态解析）| SQLite（巡检记录）\n'
    '第5层-模型层：DeepSeek API（云端）| Ollama+qwen3:8b（本地）| '
    'MiniLM-L12-v2（嵌入模型，CPU）\n'
    '第6层-基础设施：Python 3.9+ | Windows/Linux | 虚拟环境\n\n'
    '请用 Mermaid 或文字框图的形式描述上述架构，适合直接放入答辩 PPT。'
)

doc.add_heading('提示词三：生成演示脚本', level=2)

prompt3 = doc.add_paragraph()
run = prompt3.add_run(
    '请帮我写一份5分钟的项目答辩演示脚本，演示《OA运维多智能Agent巡检问答系统》的'
    '知识问答和知识库功能。\n\n'
    '演示流程：\n'
    '1. 打开系统，展示仪表盘（0.5分钟）\n'
    '2. 进入知识库，演示上传一份OA运维文档（PDF），展示解析结果和KG实体（1分钟）\n'
    '3. 进入知识问答，问一个简单问题如"OA审批流程是什么？"，展示RAG回答和命令推荐（1分钟）\n'
    '4. 追问一个关联问题如"审批卡死了怎么排查？"，展示多轮对话记忆（1分钟）\n'
    '5. 进入图谱探索，搜索实体并查看子图和路径（0.5分钟）\n'
    '6. 切换到系统设置，演示一键切换到DeepSeek云端模型，再次问答展示差异（0.5分钟）\n'
    '7. 总结技术亮点（0.5分钟）\n\n'
    '要求：脚本语言自然、有互动感，标注每步的关键话术和操作说明。'
)

doc.add_paragraph()
doc.add_paragraph('— 文档结束 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============ 保存 ============
desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
filepath = os.path.join(desktop, 'OA运维Agent_知识问答与知识库_答辩技术文档.docx')
doc.save(filepath)
print(f'文档已保存到: {filepath}')
